import os
import tempfile
from datetime import datetime

# --- DOCX generation ---
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- PDF generation (ReportLab) ---
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# --- Markdown + NLP ---
import markdown
from bs4 import BeautifulSoup
from textblob import TextBlob

def create_txt_file(content: str, filename: str = None) -> str:
    """Create a .txt file with the given content."""
    if filename is None:
        filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    
    filepath = os.path.join(tempfile.gettempdir(), filename)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✅ TXT file created: {filepath}")
    return filepath


def create_docx_file(content: str, title: str = "Generated Document", filename: str = None) -> str:
    """Create a .docx file with formatted content."""
    if filename is None:
        filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    filepath = os.path.join(tempfile.gettempdir(), filename)
    
    doc = Document()
    
    # Add title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    timestamp.runs[0].font.size = Pt(9)
    timestamp.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Blank line
    
    # Process content - handle sections and formatting
    sections = content.split('\n\n')
    
    for section in sections:
        if section.strip():
            # Check if it's a heading (starts with #, ##, etc.)
            if section.startswith('# '):
                doc.add_heading(section[2:].strip(), level=1)
            elif section.startswith('## '):
                doc.add_heading(section[3:].strip(), level=2)
            elif section.startswith('### '):
                doc.add_heading(section[4:].strip(), level=3)
            # Check if it's a code block (starts with ```)
            elif section.startswith('```'):
                code_content = section.replace('```', '').strip()
                p = doc.add_paragraph(code_content)
                p.style = 'Intense Quote'
            # Regular paragraph
            else:
                # Handle bullet points
                if section.strip().startswith('- ') or section.strip().startswith('* '):
                    for line in section.split('\n'):
                        if line.strip().startswith(('- ', '* ')):
                            doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif line.strip():
                            doc.add_paragraph(line.strip())
                # Handle numbered lists
                elif any(section.strip().startswith(f'{i}.') for i in range(1, 10)):
                    for line in section.split('\n'):
                        if line.strip() and line.strip()[0].isdigit():
                            doc.add_paragraph(line.strip().split('.', 1)[1].strip(), style='List Number')
                        elif line.strip():
                            doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph(section.strip())
    
    doc.save(filepath)
    print(f"✅ DOCX file created: {filepath}")
    return filepath


# def create_pdf_file(content: str, title: str = "Generated Document", filename: str = None) -> str:
#     """Create a .pdf file with formatted content."""
#     if filename is None:
#         filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
#     filepath = os.path.join(tempfile.gettempdir(), filename)
    
#     doc = SimpleDocTemplate(filepath, pagesize=letter)
#     styles = getSampleStyleSheet()
    
#     # Custom styles
#     title_style = ParagraphStyle(
#         'CustomTitle',
#         parent=styles['Heading1'],
#         fontSize=24,
#         textColor='#2C3E50',
#         spaceAfter=30,
#         alignment=1  # Center
#     )
    
#     heading_style = ParagraphStyle(
#         'CustomHeading',
#         parent=styles['Heading2'],
#         fontSize=16,
#         textColor='#34495E',
#         spaceAfter=12,
#         spaceBefore=12
#     )
    
#     # Build document
#     story = []
    
#     # Add title
#     story.append(Paragraph(title, title_style))
#     story.append(Spacer(1, 0.2*inch))
    
#     # Add timestamp
#     timestamp = f"<para align='right'><font size=8 color='gray'>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</font></para>"
#     story.append(Paragraph(timestamp, styles['Normal']))
#     story.append(Spacer(1, 0.3*inch))
    
#     # Process content
#     sections = content.split('\n\n')
    
#     for section in sections:
#         if section.strip():
#             # Check if it's a heading
#             if section.startswith('# '):
#                 story.append(Paragraph(section[2:].strip(), heading_style))
#             elif section.startswith('## '):
#                 story.append(Paragraph(section[3:].strip(), styles['Heading3']))
#             # Regular paragraph
#             else:
#                 # Handle special characters for PDF
#                 safe_content = section.replace('<', '&lt;').replace('>', '&gt;')
#                 # Handle bullet points
#                 if safe_content.strip().startswith(('- ', '* ')):
#                     for line in safe_content.split('\n'):
#                         if line.strip().startswith(('- ', '* ')):
#                             story.append(Paragraph(f"• {line.strip()[2:]}", styles['Normal']))
#                         elif line.strip():
#                             story.append(Paragraph(line.strip(), styles['Normal']))
#                 else:
#                     story.append(Paragraph(safe_content, styles['Normal']))
#                 story.append(Spacer(1, 0.1*inch))
    
#     doc.build(story)
#     print(f"✅ PDF file created: {filepath}")
#     return filepath




def create_pdf_file(content: str, title: str = "Generated Document", filename: str = None, summarize=False) -> str:
    """Create a .pdf file with full markdown-like formatting and optional NLP summary."""

    # Generate filename if not provided
    if filename is None:
        filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    filepath = os.path.join(tempfile.gettempdir(), filename)

    # --- NLP Summarization (Optional) ---
    if summarize:
        blob = TextBlob(content)
        sentences = blob.sentences
        if len(sentences) > 5:
            summary = " ".join(str(s) for s in sentences[:5])
            content = f"**Summary:** {summary}\n\n---\n\n{content}"

    # --- Convert Markdown → HTML ---
    html_content = markdown.markdown(
        content,
        extensions=["fenced_code", "tables", "sane_lists", "nl2br"]
    )

    # --- Setup PDF document ---
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            leftMargin=72, rightMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()

    # --- Define Custom Styles ---
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=22,
        textColor=colors.HexColor("#2C3E50"),
        alignment=1,  # Center
        spaceAfter=20
    )

    heading1_style = ParagraphStyle(
        'Heading1Custom',
        parent=styles['Heading1'],
        textColor=colors.HexColor("#1F618D"),
        fontSize=16,
        spaceAfter=10,
        spaceBefore=10
    )

    heading2_style = ParagraphStyle(
        'Heading2Custom',
        parent=styles['Heading2'],
        textColor=colors.HexColor("#2874A6"),
        fontSize=14,
        spaceAfter=8,
        spaceBefore=8
    )

    normal_style = ParagraphStyle(
        'NormalCustom',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=colors.black
    )

    code_style = ParagraphStyle(
        'CodeCustom',
        parent=styles['Code'],
        fontName='Courier',
        fontSize=9,
        backColor=colors.whitesmoke,
        textColor=colors.darkgray,
        leftIndent=12,
        rightIndent=12,
        spaceBefore=4,
        spaceAfter=4
    )

    story = []

    # --- Add Title ---
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2 * inch))

    # --- Add Timestamp ---
    timestamp = (
        f"<para align='right'><font size=8 color='gray'>"
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        f"</font></para>"
    )
    story.append(Paragraph(timestamp, styles['Normal']))
    story.append(Spacer(1, 0.3 * inch))

    # --- Parse HTML Content ---
    soup = BeautifulSoup(html_content, "html.parser")
    
    def process_element(el):
        """Recursively convert HTML to PDF elements."""
        from bs4.element import NavigableString

        # Handle plain text (no tag)
        if isinstance(el, NavigableString):
            text = str(el).strip()
            if text:
                story.append(Paragraph(text, normal_style))
            return

        if el.name in ["h1"]:
            story.append(Paragraph(el.decode_contents(), heading1_style))
        elif el.name in ["h2"]:
            story.append(Paragraph(el.decode_contents(), heading2_style))
        elif el.name in ["h3"]:
            story.append(Paragraph(el.decode_contents(), styles['Heading3']))
        elif el.name in ["ul"]:
            for li in el.find_all("li", recursive=False):
                story.append(Paragraph(f"• {li.decode_contents()}", normal_style))
        elif el.name in ["ol"]:
            for i, li in enumerate(el.find_all("li", recursive=False), 1):
                story.append(Paragraph(f"{i}. {li.decode_contents()}", normal_style))
        elif el.name == "pre" or (el.name == "code" and el.parent.name == "pre"):
            code_text = el.get_text().replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(f"<font face='Courier'>{code_text}</font>", code_style))
        elif el.name == "p" or el.name is None:
            text = el.decode_contents()
            story.append(Paragraph(text, normal_style))

        # Recurse into children
        for child in el.children:
            if hasattr(child, "name") or isinstance(child, NavigableString):
                process_element(child)



    for element in soup.contents:
        process_element(element)
        story.append(Spacer(1, 0.1 * inch))

    # --- Build PDF ---
    doc.build(story)

    print(f"✅ Enhanced PDF created successfully: {filepath}")
    return filepath



def generate_document(content: str, format: str = "txt", title: str = "Generated Document", filename: str = None) -> str:
    """
    Main function to generate documents in different formats.
    
    Args:
        content: The text content to be included in the document
        format: 'txt', 'docx', or 'pdf'
        title: Document title (used for DOCX and PDF)
        filename: Custom filename (optional)
    
    Returns:
        filepath: Path to the generated file
    """
    format = format.lower().strip()
    
    if format == 'txt' or format == '.txt':
        return create_txt_file(content, filename)
    elif format == 'docx' or format == '.docx':
        return create_docx_file(content, title, filename)
    elif format == 'pdf' or format == '.pdf':
        return create_pdf_file(content, title, filename)
    else:
        raise ValueError(f"Unsupported format: {format}. Use 'txt', 'docx', or 'pdf'")